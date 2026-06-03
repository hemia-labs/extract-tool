from pathlib import Path

from docx import Document

from app.extractors.base import BaseExtractor, ExtractionResult


class DocxExtractor(BaseExtractor):
    def extract(self, path: Path, ocr: str, language: str, output: str) -> ExtractionResult:
        document = Document(path)
        parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        text = "\n".join(parts)
        return {
            "text": text,
            "pages": [{"page": 1, "text": text, "confidence": None}],
            "metadata": {"ocrUsed": False, "pages": 1, "characters": len(text)},
        }
